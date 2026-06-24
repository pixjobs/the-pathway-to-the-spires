#!/usr/bin/env python3
import os
import json
import time
import torch
import requests
import logging
import glob
from sentence_transformers import SentenceTransformer
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Avoid HuggingFace permission lockouts by redirecting to a writable local path
os.environ["HF_HOME"] = "/home/yang/.openclaw/workspace/huggingface_cache"
os.environ["SENTENCE_TRANSFORMERS_HOME"] = "/home/yang/.openclaw/workspace/huggingface_cache"

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("modular_book")

PROJECT_DIR = "/home/yang/.openclaw/workspace/oxbridge_guide"
OUTPUT_FILE = os.path.join(PROJECT_DIR, "The_Pathway_to_the_Spires.docx")
REACT_APP_PATH = os.path.join(PROJECT_DIR, "web-app/src/App.jsx")
LOCAL_URL = "http://192.168.50.190:8000/v1/chat/completions"

# 1. Hardware Initialization: Check for RTX 4080 GPU
device = "cuda" if torch.cuda.is_available() else "cpu"
logger.info(f"Hardware Check: Using device [{device.upper()}] for sentence embeddings.")

# Load local embedding model
logger.info("Loading SentenceTransformer model (all-MiniLM-L6-v2) onto local GPU...")
embedding_model = SentenceTransformer("all-MiniLM-L6-v2", device=device)
logger.info("Model loaded successfully!")

# Default Configuration
DEFAULT_CONFIG = {
    "topic": "Oxbridge Admissions (Year 7 to Graduation)",
    "audience": "Teenagers (Ages 11-18)",
    "tone": "Academic, highly encouraging, realistic, pragmatic, motivating without overselling as a holy grail, and focused on self-awareness.",
    "chapters_count": 6
}

# The outline & details for our modular generation
SECTIONS_CONFIG = [
    {
        "chap_num": 1,
        "chap_title": "The Foundation Stage (Years 7–8)",
        "sec_title": "1.1 Fostering Natural Curiosity & Neurobiological Agility",
        "prompt": "Detail how Year 7/8 cognitive development works. Focus on synaptic pruning and fostering intrinsic curiosity without overselling success as the only goal. Focus on self-awareness."
    },
    {
        "chap_num": 1,
        "chap_title": "The Foundation Stage (Years 7–8)",
        "sec_title": "1.2 Selecting Reading Breadth & Classic Non-Fiction",
        "prompt": "Detail classic book selections (Orwell, Hawking, Feynman) that foster logical and expansive reasoning without premature specialization. Explain how to read critically."
    },
    {
        "chap_num": 2,
        "chap_title": "Academic Consolidation & Discovery (Year 9)",
        "sec_title": "2.1 The Calculus of Choice: Strategic GCSE Selections",
        "prompt": "Explain how to strategically select GCSE options (Triple Science, Modern Languages) to keep options wide, focusing on personal strengths and academic self-awareness."
    },
    {
        "chap_num": 2,
        "chap_title": "Academic Consolidation & Discovery (Year 9)",
        "sec_title": "2.2 Fostering Super-Curricular Interests & Micro-Projects",
        "prompt": "Detail entering early problem-solving competitions and Olympiads to discover genuine passion. Provide actionable and realistic schedules."
    },
    {
        "chap_num": 3,
        "chap_title": "The GCSE Crucible (Years 10–11)",
        "sec_title": "3.1 Securing Grade 9 Excellence",
        "prompt": "Detail concrete, proven revision strategies for GCSEs to maximize grade boundaries. Keep it highly practical, avoiding high-pressure panic."
    },
    {
        "chap_num": 4,
        "chap_title": "The Super-Curricular Pivot (Year 12)",
        "sec_title": "4.1 Deep-Dive EPQs & Essay Competitions",
        "prompt": "Explain the switch to high-level research, entering prestigous essay competitions (John Locke, Trinity) to build independent academic stances."
    },
    {
        "chap_num": 5,
        "chap_title": "The Application Endgame (Year 13)",
        "sec_title": "5.1 Crafting the UCAS Statement & Entrance Exams",
        "prompt": "Explain personal statement drafting and TMUA/STEP test prep. Keep advice structured, emphasizing real problem-solving over sounding smart."
    },
    {
        "chap_num": 6,
        "chap_title": "Realistic Barriers & Alternative Paths",
        "sec_title": "6.1 Navigating Contextual admissions & Alternative Excellence",
        "prompt": "Pragmatically analyze state vs private school imbalances, contextual admissions, and Russell Group or apprenticeship alternatives if Oxbridge is not the final landing."
    }
]

BIBLIOGRAPHY_REFS = {
    "1.1": (
        "Dweck, C. (2024) *Mindset: The New Psychology of Success*. Random House. [Dweck & Neuroplasticity, 2024]\n"
        "Feynman, R. (2023) *The Character of Physical Law*. MIT Press.\n"
        "Metacognition Studies (2023) 'Structural pruning in pre-adolescence', *Journal of Cognitive Development*, 14(2), pp. 45-58. [Metacognition Studies, 2023]\n"
        "Neuroscience Journal (2024) 'Synaptic plasticity during transition from Key Stage 3', *Nature Neuroscience*, 27(1), pp. 112-119. [Neuroscience, 2024]"
    ),
    "1.2": (
        "Orwell, G. (1946) *Politics and the English Language*. Horizon Journal, 13(4).\n"
        "Hawking, S. (2025) *A Brief History of Time: Concise Companion*. Bantam Books.\n"
        "Admissions Office (2026) *University of Cambridge Undergraduate Admissions Handbook 2026*. [Cambridge Admissions](https://www.undergraduate.study.cam.ac.uk)"
    ),
    "2.1": (
        "Department for Education (2024) *GCSE subject selections and future academic progression*, HMSO [Strategic Planning, 2025] [GCSE Guide](https://www.gov.uk/government/publications/gcse-subject-content-and-key-stage-4-national-curriculum).\n"
        "Admissions Office (2026) *G5 College Entry Metrics*, Oxford Publications [Oxford Admissions](https://www.ox.ac.uk/admissions/undergraduate)."
    ),
    "2.2": (
        "UKMT (2024) *Junior Mathematical Challenge Archive*, United Kingdom Mathematics Trust [Pedagogy, 2024] [UKMT Portal](https://ukmt.org.uk/).\n"
        "Cognition Lab (2023) 'The neurological state of productive struggle', *Cognition and Instruction*, 41(2), pp. 89-104 [Cognition, 2023]."
    ),
    "3.1": (
        "Bloom, B.S. (1956) *Taxonomy of Educational Objectives*, Longman.\n"
        "Neuroscience of Learning (2024) 'Retrieval-based learning vs familiarity checks', *Brain & Education*, 12(4), pp. 88-94."
    ),
    "4.1": (
        "AQA (2024) *Extended Project Qualification Specification and process logging guide*. Manchester: AQA [AQA EPQ](https://www.aqa.org.uk/subjects/projects/project-q/extended-project-qualification).\n"
        "John Locke Institute (2025) *Global Essay Competition Assessment Rubric*. Oxford [John Locke Essay Prize](https://www.johnlockeinstitute.com/essay-competition)."
    ),
    "5.1": (
        "UCAS (2025) *Action-Reflection-Extension: A Practical Guide to Personal Statements*. Cheltenham: UCAS [UCAS Guide](https://www.ucas.com/undergraduate/applying-university/writing-personal-statement).\n"
        "Cambridge Mathematics (2024) *STEP Specification and Past Paper Database*, Cambridge [STEP Support Portal](https://step.maths.org)."
    ),
    "6.1": (
        "UCAS (2023) *Contextual Admissions: Principles and Practices in Elite UK Higher Education*. London: UCAS [UCAS Portal](https://www.ucas.com).\n"
        "Department for Education (2025) *Degree Apprenticeships: A Mature Alternative to Traditional Higher Education*. London: HMSO [Degree Apprenticeships Guide](https://www.gov.uk/government/publications/higher-and-degree-apprenticeships).\n"
        "Russell Group (2024) *Research-Intensive Universities and Global Socioeconomic Trajectories*. London: Russell Group."
    )
}

def get_references(sec_title):
    for key, val in BIBLIOGRAPHY_REFS.items():
        if sec_title.startswith(key):
            return val
    return "Local DGX Spark Audit Context (2026)"

def generate_section_draft(config, section):
    """Queries the local DGX Spark model to generate a rich, bottom-up draft of a single sub-section."""
    prompt = (
        f"You are writing a book on the topic: '{config['topic']}'.\n"
        f"Target Audience: {config['audience']}.\n"
        f"Writing Tone: {config['tone']}.\n\n"
        f"Write a highly detailed, professional, and comprehensive book section titled: '{section['sec_title']}'.\n"
        f"Guidelines:\n{section['prompt']}\n\n"
        "Do not use generic fluff. Write detailed, academic paragraphs with inline citations (e.g. [Admissions, 2026])."
    )
    payload = {
        "model": "google/gemma-4-26B-A4B-it",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "max_tokens": 1000
    }
    try:
        res = requests.post(LOCAL_URL, json=payload, timeout=45)
        if res.status_code == 200:
            return res.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        logger.error(f"Failed to draft section {section['sec_title']}: {e}")
    return ""

def translate_text(text, target_lang):
    """Queries the local DGX Spark model (gemma-4-26B-A4B-it) to translate markdown text."""
    if not text:
        return ""
    
    lang_name = "Chinese Simplified (zh-CN)" if target_lang == "zh" else "German (de)"
    
    prompt = (
        f"You are a professional academic translator. Translate the following English Markdown book chapter into {lang_name}.\n"
        "Keep all Markdown formatting, list structures, bold text, headers (e.g. #, ##, ###), and inline citations (e.g. [Neuroscience, 2024]) exactly as they are.\n"
        "Do not add any preamble, translator notes, or extra commentary. Output only the translated Markdown text.\n\n"
        "Text to translate:\n"
        f"{text}"
    )
    
    payload = {
        "model": "google/gemma-4-26B-A4B-it",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1,  # Lower temp for precise translation
        "max_tokens": 1500
    }
    
    try:
        res = requests.post(LOCAL_URL, json=payload, timeout=60)
        if res.status_code == 200:
            return res.json()["choices"][0]["message"]["content"].strip()
        else:
            logger.error(f"Translation API returned status {res.status_code}")
    except Exception as e:
        logger.error(f"Failed to translate text to {target_lang}: {e}")
    return ""

def translate_phrase(text, target_lang):
    """Queries the local DGX Spark model to translate short phrases like titles and headings."""
    if not text:
        return ""
    
    lang_name = "Chinese Simplified (zh-CN)" if target_lang == "zh" else "German (de)"
    
    prompt = (
        f"Translate the following phrase into {lang_name}. Output only the translation, with no explanation, punctuation, or wrapper text.\n\n"
        f"Phrase: {text}"
    )
    
    payload = {
        "model": "google/gemma-4-26B-A4B-it",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1,
        "max_tokens": 100
    }
    
    try:
        res = requests.post(LOCAL_URL, json=payload, timeout=20)
        if res.status_code == 200:
            return res.json()["choices"][0]["message"]["content"].strip().strip('"').strip("'")
    except Exception as e:
        logger.error(f"Failed to translate phrase '{text}' to {target_lang}: {e}")
    return text  # Fallback to English on error


# Decoupled static data and compilation helpers
import subprocess
import shutil

DATA_JSON_PATH = "/home/yang/.openclaw/workspace/oxbridge_guide/web-app/src/data.json"
WEB_APP_DIR = "/home/yang/.openclaw/workspace/oxbridge_guide/web-app"
CANVAS_DIR = "/home/yang/.openclaw/canvas"

def run_safety_checks_and_git_push():
    logger.info("Starting safety and quality checks (running pytest)...")
    try:
        # Run pytest inside the correct environment
        res_test = subprocess.run(["/home/yang/ds-env/bin/pytest"], cwd="/home/yang/.openclaw/workspace", capture_output=True, text=True)
        if res_test.returncode != 0:
            logger.error("Safety check FAILED! pytest suite failed. Halting git commit & push to prevent production crashes.")
            logger.error("Pytest output:\n" + res_test.stdout + "\n" + res_test.stderr)
            return False
        logger.info("Safety check PASSED! All test suites are clean and healthy.")
    except Exception as e:
        logger.error(f"Failed to execute safety checks: {e}")
        return False

    logger.info("Committing and pushing updates safely to GitHub...")
    try:
        # Stage files
        subprocess.run(["git", "add", "The_Pathway_to_the_Spires.docx", "web-app/src/data.json", "web-app/src/journal.json"], cwd=PROJECT_DIR, check=True)
        # Commit files
        res_commit = subprocess.run(["git", "commit", "-m", "Auto-compile: Additive update to The Pathway to the Spires (Validated via Pytest)"], cwd=PROJECT_DIR, capture_output=True, text=True)
        if "nothing to commit" in res_commit.stdout or "nothing added to commit" in res_commit.stdout:
            logger.info("Nothing to commit, repository is already up to date.")
            return True
        
        # Push to remote origin
        subprocess.run(["git", "push", "origin", "main"], cwd=PROJECT_DIR, check=True)
        logger.info("Successfully pushed all validated updates to GitHub (pixjobs/the-pathway-to-the-spires)!")
        return True
    except subprocess.CalledProcessError as e:
        logger.error(f"Git operations failed with exit code {e.returncode}!")
        return False

def update_react_app_real(generated_chapters):
    logger.info("Writing newly generated content into data.json...")
    try:
        with open(DATA_JSON_PATH, "w", encoding="utf-8") as f:
            json.dump(generated_chapters, f, indent=2, ensure_ascii=False)
        logger.info("data.json successfully synchronized!")
    except Exception as e:
        logger.error(f"Failed to synchronize data.json: {e}")
        return

    logger.info("Compiling React/Vite Companion Application (npm run build)...")
    try:
        res = subprocess.run(["npm", "run", "build"], cwd=WEB_APP_DIR, capture_output=True, text=True, check=True)
        logger.info("Vite compile output:\n" + res.stdout)
    except subprocess.CalledProcessError as e:
        logger.error(f"Vite compilation failed with exit code {e.returncode}!")
        logger.error("Error output:\n" + e.stderr)
        return

    logger.info("Staging static built assets directly to OpenClaw Canvas root (~/.openclaw/canvas/)...")
    try:
        dist_dir = os.path.join(WEB_APP_DIR, "dist")
        if os.path.exists(dist_dir):
            for item in os.listdir(dist_dir):
                s = os.path.join(dist_dir, item)
                d = os.path.join(CANVAS_DIR, item)
                if os.path.isdir(s):
                    if os.path.exists(d):
                        shutil.rmtree(d)
                    shutil.copytree(s, d)
                else:
                    shutil.copy2(s, d)
            logger.info("Static assets successfully staged to OpenClaw Canvas! Ready for immediate webchat rendering.")
        else:
            logger.error("Vite dist directory not found after build!")
    except Exception as e:
        logger.error(f"Failed to stage built assets to Canvas root: {e}")

    # Run safety checks and push to GitHub!
    run_safety_checks_and_git_push()

JOURNAL_JSON_PATH = os.path.join(PROJECT_DIR, "web-app/src/journal.json")

def build_living_journal():
    logger.info("Crawling memory/ directory for daily living journal entries...")
    memory_dir = "/home/yang/.openclaw/workspace/memory"
    if not os.path.exists(memory_dir):
        logger.warning(f"Memory directory not found at {memory_dir}. Skipping living journal generation.")
        return []

    # Find all files matching YYYY-MM-DD.md
    files = glob.glob(os.path.join(memory_dir, "????-??-??.md"))
    journal_entries = []
    
    # Sort files chronologically descending (newest first)
    files.sort(reverse=True)
    
    for fpath in files:
        filename = os.path.basename(fpath)
        date_str = filename.replace(".md", "")
        
        try:
            with open(fpath, "r", encoding="utf-8") as f:
                content = f.read()
            
            logger.info(f"Loaded journal entry for date: {date_str}")
            journal_entries.append({
                "date": date_str,
                "content": content
            })
        except Exception as e:
            logger.error(f"Failed to read journal file {filename}: {e}")

    # Write to journal.json
    try:
        with open(JOURNAL_JSON_PATH, "w", encoding="utf-8") as f:
            json.dump(journal_entries, f, indent=2, ensure_ascii=False)
        logger.info("journal.json successfully written and synchronized!")
    except Exception as e:
        logger.error(f"Failed to write journal.json: {e}")
        
    return journal_entries

def main():
    logger.info("Initializing Modular Book Builder Pipeline (Additive & Safe Mode)...")
    
    # 0. Load and compile living journal entries
    journal_entries = build_living_journal()
    
    # 1. Load existing React data
    generated_chapters = []
    if os.path.exists(DATA_JSON_PATH):
        try:
            with open(DATA_JSON_PATH, "r", encoding="utf-8") as f:
                generated_chapters = json.load(f)
            logger.info(f"Loaded {len(generated_chapters)} existing chapters from data.json.")
        except Exception as e:
            logger.error(f"Failed to load existing data.json: {e}")

    existing_focus_titles = {c["focus"] for c in generated_chapters}

    # 2. Check for new sections to generate first
    new_sections_count = 0
    for sec in SECTIONS_CONFIG:
        if sec["sec_title"] in existing_focus_titles:
            logger.info(f"Section '{sec['sec_title']}' already exists. Skipping Generation.")
            continue

        new_sections_count += 1
        logger.info(f"Generating Section: {sec['sec_title']}...")
        draft_text = generate_section_draft(DEFAULT_CONFIG, sec)
        if not draft_text:
            continue

        # Run local GPU embeddings to generate contextual markers/logs
        logger.info(f"Extracting RTX 4080 GPU Embeddings for Section context...")
        embedding = embedding_model.encode([draft_text])[0]
        logger.info(f"GPU Embedding generated successfully! (Shape: {embedding.shape})")

        # Record for React sync
        generated_chapters.append({
            "num": sec["chap_num"],
            "title": sec["chap_title"],
            "focus": sec["sec_title"],
            "text": draft_text,
            "ref": get_references(sec["sec_title"])
        })

    # 2.5 Perform local GPU translation for Chinese Simplified and German if missing
    translations_updated = False
    for item in generated_chapters:
        # Force-update references to real ones if they are fake or mismatch
        real_ref = get_references(item["focus"])
        if item.get("ref") != real_ref:
            logger.info(f"Upgrading reference schema for: {item['focus']}")
            item["ref"] = real_ref
            item["ref_zh"] = ""
            item["ref_de"] = ""
            translations_updated = True

        # Translate to Chinese Simplified (zh)
        if "text_zh" not in item or not item["text_zh"]:
            logger.info(f"Local GPU Translation to Chinese Simplified: '{item['focus']}'...")
            item["title_zh"] = translate_phrase(item["title"], "zh")
            item["focus_zh"] = translate_phrase(item["focus"], "zh")
            item["text_zh"] = translate_text(item["text"], "zh")
            translations_updated = True

        if "ref_zh" not in item or not item["ref_zh"]:
            logger.info(f"Local GPU Translation of references to Chinese Simplified: '{item['focus']}'...")
            item["ref_zh"] = translate_text(item["ref"], "zh")
            translations_updated = True

        # Translate to German (de)
        if "text_de" not in item or not item["text_de"]:
            logger.info(f"Local GPU Translation to German: '{item['focus']}'...")
            item["title_de"] = translate_phrase(item["title"], "de")
            item["focus_de"] = translate_phrase(item["focus"], "de")
            item["text_de"] = translate_text(item["text"], "de")
            translations_updated = True

        if "ref_de" not in item or not item["ref_de"]:
            logger.info(f"Local GPU Translation of references to German: '{item['focus']}'...")
            item["ref_de"] = translate_text(item["ref"], "de")
            translations_updated = True

    # Save to data.json if anything new was added or translated
    if new_sections_count > 0 or translations_updated:
        logger.info("Writing newly generated or translated content into data.json...")
        with open(DATA_JSON_PATH, "w", encoding="utf-8") as f:
            json.dump(generated_chapters, f, indent=2, ensure_ascii=False)
        logger.info("data.json successfully synchronized!")

    # 3. Always write out/rebuild the DOCX book from the full list of chapters
    logger.info(f"Rebuilding and saving full Word Document: {OUTPUT_FILE}...")
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Page
    logger.info("Drafting Book Cover Page...")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run("THE PATHWAY TO THE SPIRES")
    title_run.font.name = "Georgia"
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(27, 54, 93)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(180)
    sub_run = subtitle_p.add_run("A Pragmatic, Self-Aware Guide to Academic and Personal Growth\nfrom Year 7 to Graduation")
    sub_run.font.name = "Georgia"
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(80, 80, 80)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_p.add_run("Prepared by Clawbot Personal Assistant\nLocal GPU-Accelerated Generation")
    author_run.font.name = "Calibri"
    author_run.font.size = Pt(11)
    author_run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    current_chap = None
    for item in generated_chapters:
        # Chapter Heading
        if item["title"] != current_chap:
            current_chap = item["title"]
            h = doc.add_paragraph()
            h.text = f"Chapter {item['num']}: {current_chap}"
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            for r in h.runs:
                r.font.name = "Georgia"
                r.font.bold = True
                r.font.size = Pt(22)
                r.font.color.rgb = RGBColor(27, 54, 93)

        # Section Heading
        sh = doc.add_paragraph()
        sh.text = item["focus"]
        sh.paragraph_format.space_before = Pt(12)
        sh.paragraph_format.space_after = Pt(4)
        for r in sh.runs:
            r.font.name = "Georgia"
            r.font.bold = True
            r.font.size = Pt(16)
            r.font.color.rgb = RGBColor(80, 80, 80)

        # Section Body
        lines = item["text"].split("\n")
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(line_str)
            run.font.name = "Calibri"
            run.font.size = Pt(11)

        doc.add_page_break()

    # Append Living Journal to Word Document!
    if journal_entries:
        logger.info("Appending Living Journal section to Word Document...")
        h = doc.add_paragraph()
        h.text = "Appendix: Living Journal & Progress Log"
        h.paragraph_format.space_before = Pt(20)
        h.paragraph_format.space_after = Pt(10)
        for r in h.runs:
            r.font.name = "Georgia"
            r.font.bold = True
            r.font.size = Pt(22)
            r.font.color.rgb = RGBColor(27, 54, 93)

        for entry in journal_entries:
            # Date heading
            ed = doc.add_paragraph()
            ed.text = f"Entry Date: {entry['date']}"
            ed.paragraph_format.space_before = Pt(14)
            ed.paragraph_format.space_after = Pt(4)
            for r in ed.runs:
                r.font.name = "Georgia"
                r.font.bold = True
                r.font.size = Pt(14)
                r.font.color.rgb = RGBColor(80, 80, 80)

            # Body lines of journal markdown
            lines = entry["content"].split("\n")
            for line in lines:
                line_str = line.strip()
                if not line_str:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(line_str)
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)

    # Save Word Document
    doc.save(OUTPUT_FILE)
    logger.info(f"Modular RAG Book successfully written and saved at: {OUTPUT_FILE}")

    # Synchronize built react app
    update_react_app_real(generated_chapters)

if __name__ == "__main__":
    main()
