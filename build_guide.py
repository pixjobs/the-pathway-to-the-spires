#!/usr/bin/env python3
import os
import json
import time
import requests
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("oxbridge_builder")

PROJECT_DIR = "/home/yang/.openclaw/workspace/oxbridge_guide"
OUTPUT_FILE = os.path.join(PROJECT_DIR, "The_Pathway_to_the_Spires.docx")

# Load Gemini API key from openclaw.json
openclaw_config = "/home/yang/.openclaw/openclaw.json"
gemini_key = ""
if os.path.exists(openclaw_config):
    try:
        with open(openclaw_config, "r") as f:
            cfg = json.load(f)
        gemini_key = cfg.get("models", {}).get("providers", {}).get("google", {}).get("apiKey", "")
    except Exception as e:
        logger.error(f"Failed to load OpenClaw config: {e}")

GEMINI_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={gemini_key}"

CHAPTERS = [
    {
        "num": 1,
        "title": "The Foundation Stage (Years 7–8)",
        "prompt": "Write a massive, academic-grade Chapter 1 (at least 1200 words) focusing on Year 7 and 8 cognitive development, reading breadth, logical reasoning, and establishing excellent study habits. Detail how to foster curiosity, select academic books (mention classic titles like Orwell, Hawking, and Feynman), and introduce logical problem solving. Do not specialize too early. Keep it highly detailed, professional, and include inline academic citations and a reference list at the end."
    },
    {
        "num": 2,
        "title": "Academic Consolidation & Discovery (Year 9)",
        "prompt": "Write a massive, academic-grade Chapter 2 (at least 1200 words) focusing on Year 9. Detail strategic GCSE option selections (Triple Science, Modern Languages, Humanities), building super-curricular interests, entering competitive mathematical/scientific or literary problem-solving, and establishing self-directed learning paths. Keep it highly realistic, professional, and include inline academic citations and a reference list at the end."
    },
    {
        "num": 3,
        "title": "The GCSE Crucible (Years 10–11)",
        "prompt": "Write a massive, academic-grade Chapter 3 (at least 1200 words) focusing on Years 10 and 11. Detail how to maximize Grade 9 achievements, participating in UK regional/national Olympiads (like UKMT, Chemistry/Physics Olympiads), and strategically selecting A-Level or IB choices (matching university course requirements). Keep it academic, highly structured, and include inline academic citations and a reference list at the end."
    },
    {
        "num": 4,
        "title": "The Super-Curricular Pivot (Year 12)",
        "prompt": "Write a massive, academic-grade Chapter 4 (at least 1200 words) focusing on Year 12. Detail the crucial shift to super-curricular deep-dives: EPQs, university-level textbooks, entering prestigious essay competitions (like Trinity College, John Locke, Marshall), and initial prep for admissions tests (TMUA, STEP, MAT, HAT, PAT, TSA, LNAT). Keep it highly detailed, precise, and include inline academic citations and a reference list at the end."
    },
    {
        "num": 5,
        "title": "The Application, Test, & Interview Endgame (Year 13)",
        "prompt": "Write a massive, academic-grade Chapter 5 (at least 1200 words) focusing on Year 13. Detail personal statement structuring, intensive admissions test preparation, Oxbridge mock interview simulations (focusing on the Oxbridge tutorial style of oral problem solving), and navigating the December interview pools. Keep it realistic, academic, and include inline academic citations and a reference list at the end."
    },
    {
        "num": 6,
        "title": "Realistic Barriers & Alternative Pathways",
        "prompt": "Write a massive, academic-grade Chapter 6 (at least 1200 words) focusing on realistic socioeconomic or school-type barriers, contextual admissions mitigation strategies, and high-quality alternative options (like Russell Group universities, US Ivy Leagues, or degree apprenticeships) if Oxbridge is not the final landing. Keep it pragmatic, encouraging, and include inline academic citations and a reference list at the end."
    }
]

def query_gemini_generate(prompt, chapter_title):
    sys_instruction = (
        "You are an Oxford and Cambridge University admissions coach, professional editor, and academic researcher. "
        "Your style is elegant, intellectual, encouraging, and rich in technical details, perfect for a holiday reading guide. "
        "Write extremely comprehensive, detailed, and expansive chapters (each at least 1200 words) with inline citations and robust academic reasoning."
    )
    payload = {
        "contents": [{"parts": [{"text": f"{sys_instruction}\n\nTask: {prompt}"}]}],
        "generationConfig": {
            "maxOutputTokens": 4096,
            "temperature": 0.3
        }
    }
    headers = {"Content-Type": "application/json"}
    try:
        res = requests.post(GEMINI_URL, headers=headers, json=payload, timeout=60)
        if res.status_code == 200:
            return res.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
        else:
            logger.error(f"Gemini API returned error {res.status_code}: {res.text}")
    except Exception as e:
        logger.error(f"Gemini generation failed: {e}")
    return ""

def style_heading(paragraph, text, level, font_name="Georgia", color=RGBColor(27, 54, 93)):
    paragraph.text = text
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.bold = True
        run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(22)
        elif level == 2:
            run.font.size = Pt(16)
        else:
            run.font.size = Pt(12)

def main():
    logger.info("Starting High-Capacity Oxbridge Guide Book Generator...")
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 1. Create a beautiful Cover Page
    logger.info("Generating Cover Page...")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run("THE PATHWAY TO THE SPIRES")
    title_run.font.name = "Georgia"
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(27, 54, 93)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(180)
    sub_run = subtitle_p.add_run("A Complete, Comprehensive Guide to Oxbridge Admissions\nfrom Year 7 to Year 13 Graduation")
    sub_run.font.name = "Georgia"
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(80, 80, 80)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_p.add_run("Prepared by Clawbot Personal Assistant\nJune 2026")
    author_run.font.name = "Calibri"
    author_run.font.size = Pt(11)
    author_run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # 2. Generate Chapters sequentially
    for chap in CHAPTERS:
        title = f"Chapter {chap['num']}: {chap['title']}"
        logger.info(f"Generating: {title}...")
        
        # Step A: Direct Deep Generation
        content_text = query_gemini_generate(chap["prompt"], title)
        if not content_text:
            logger.warning(f"Failed to generate {title}; skipping.")
            continue
            
        # Step B: Write Chapter Heading
        h = doc.add_paragraph()
        style_heading(h, title, level=1)
        
        # Parse polished markdown headings/paragraphs and add to document
        lines = content_text.split("\n")
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            
            if line_str.startswith("### "):
                p = doc.add_paragraph()
                style_heading(p, line_str[4:], level=3, font_name="Calibri", color=RGBColor(80, 80, 80))
            elif line_str.startswith("## "):
                p = doc.add_paragraph()
                style_heading(p, line_str[3:], level=2, font_name="Georgia", color=RGBColor(27, 54, 93))
            elif line_str.startswith("# "):
                # Skip secondary Chapter headings if redundant, otherwise style as sub-heading
                p = doc.add_paragraph()
                style_heading(p, line_str[2:], level=2, font_name="Georgia", color=RGBColor(27, 54, 93))
            elif line_str.startswith("* ") or line_str.startswith("- "):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(line_str[2:])
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                # Simple markdown bold parse (**bold**)
                parts = line_str.split("**")
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                    if idx % 2 == 1:
                        run.font.bold = True
                        
        doc.add_page_break()

    # Save document
    doc.save(OUTPUT_FILE)
    logger.info(f"Book successfully compiled and saved at: {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
